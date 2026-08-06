from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


src = Path(sys.argv[1])
dst = Path(sys.argv[2])
doc = Document(src)

new_abstract = (
    "本设计面向H题车载平衡滚球运动控制系统，构建七路红外循迹底盘、动态视觉测量、串级平衡"
    "控制和分路电源系统。视觉端以动态端点重建800×64条带，融合空槽差分、暗体特征与"
    "Alpha-Beta跟踪，50 Hz输出钢球位置；平衡端结合位置PID、BMI088加速度前馈及M2006位置—"
    "速度双环；底盘采用四轮两驱差速结构和七路加权质心循迹。视觉软件211项测试全部通过，"
    "完整链路55～58帧/s，参考匹配素材空槽527帧零候选、带球975帧候选覆盖100%，理论单圈约"
    "16.7 s。系统架构、接口及软件已验证，钢球误差、停车偏差和整车圈时须以最终场地数据验收。"
)

abstract_heading = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "摘  要")
paragraph = doc.paragraphs[abstract_heading + 1]
for run in list(paragraph.runs):
    paragraph._p.remove(run._r)
run = paragraph.add_run(new_abstract)
run.font.name = "Times New Roman"
run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
run.font.size = Pt(12)

doc.core_properties.author = ""
doc.core_properties.last_modified_by = ""
doc.core_properties.comments = ""
doc.save(dst)
