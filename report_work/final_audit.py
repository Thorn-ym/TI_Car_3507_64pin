from __future__ import annotations

import json
import re
import sys
import zipfile
from pathlib import Path

import pdfplumber
from docx import Document
from lxml import etree


docx_path = Path(sys.argv[1])
pdf_path = Path(sys.argv[2])
doc = Document(docx_path)

paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
all_text = "\n".join(paragraphs)
table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
all_content = all_text + "\n" + table_text

expected_h1 = [
    "一、系统方案",
    "二、系统理论分析与计算",
    "三、电路与程序设计",
    "四、测试方案与测试结果",
    "五、参考文献",
    "附录1：电路原理图",
    "附录2：源程序",
]

heading_counts = {heading: paragraphs.count(heading) for heading in expected_h1}
anonymous_patterns = [r"院校", r"指导教师", r"学号", r"队员姓名", r"模块负责人", r"分工", r"第一部分", r"第二部分"]
anonymous_hits = {pat: re.findall(pat, all_content) for pat in anonymous_patterns if re.search(pat, all_content)}

abstract = ""
if "摘  要" in paragraphs and any(p.startswith("关键词：") for p in paragraphs):
    start = paragraphs.index("摘  要") + 1
    end = next(i for i, p in enumerate(paragraphs) if p.startswith("关键词："))
    abstract = "".join(paragraphs[start:end])

keyword_line = next((p for p in paragraphs if p.startswith("关键词：")), "")
keywords = [x for x in keyword_line.replace("关键词：", "").split("；") if x]

with zipfile.ZipFile(docx_path) as zf:
    document_xml = zf.read("word/document.xml")
    root = etree.fromstring(document_xml)
    ns = {"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"}
    equations = len(root.xpath(".//m:oMath", namespaces=ns))
    names = set(zf.namelist())
    core = zf.read("docProps/core.xml").decode("utf-8", errors="ignore") if "docProps/core.xml" in names else ""
    custom_props = "docProps/custom.xml" in names

page_texts = []
with pdfplumber.open(pdf_path) as pdf:
    for page in pdf.pages:
        page_texts.append(page.extract_text() or "")

body_start = next(i for i, text in enumerate(page_texts, 1) if "一、系统方案" in text)
ref_start = next(i for i, text in enumerate(page_texts, 1) if "五、参考文献" in text)
body_pages = ref_start - body_start

report = {
    "docx": str(docx_path),
    "pdf_pages": len(page_texts),
    "body_page_start": body_start,
    "reference_page_start": ref_start,
    "body_pages": body_pages,
    "heading_counts": heading_counts,
    "abstract_chars": len(abstract),
    "keywords": keywords,
    "equation_objects": equations,
    "tables": len(doc.tables),
    "sections": len(doc.sections),
    "anonymous_hits": anonymous_hits,
    "custom_properties_present": custom_props,
    "core_creator_empty": "<dc:creator/>" in core or "<dc:creator></dc:creator>" in core,
    "core_last_modified_by_empty": "<cp:lastModifiedBy/>" in core or "<cp:lastModifiedBy></cp:lastModifiedBy>" in core,
    "has_placeholder_xxx": "XXX" in all_content,
    "has_code_fence": "```" in all_content,
    "page_sizes": sorted({(round(s.page_width.cm, 2), round(s.page_height.cm, 2)) for s in doc.sections}),
    "margins": sorted({(round(s.top_margin.cm, 2), round(s.bottom_margin.cm, 2), round(s.left_margin.cm, 2), round(s.right_margin.cm, 2)) for s in doc.sections}),
}
print(json.dumps(report, ensure_ascii=False, indent=2))
