from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


src = Path(sys.argv[1])
dst = Path(sys.argv[2])
doc = Document(src)

old = (
    "综合现有证据，本设计已完成H题所需的系统架构、红外循迹、视觉测量、滚球控制、计时显示"
    "和安全保护设计，模块软件与实时性指标达到预期，并具备完成赛题的理论条件；但由于缺少"
    "静态滚球误差、AB段时间、整圈钢球误差、停车偏差及无线图传实机数据，现阶段不能正式判定"
    "整机已全部达到题目要求。补齐表4所列实测项目且最差值满足限值后，方可作出“本设计达到"
    "H题全部要求”的最终验收结论。"
)
new = (
    "综合现有证据，系统已完成红外循迹、视觉测量、滚球控制、计时显示和安全保护设计，软件与"
    "实时性达到预期并具备完成赛题的理论条件；但静态滚球误差、AB段时间、整圈球位误差、停车"
    "偏差和无线图传仍缺实测数据，故暂不能判定整机全部达标。上述项目的最差值满足题目限值后，"
    "方可形成最终达标结论。"
)

found = False
for paragraph in doc.paragraphs:
    if paragraph.text == old:
        for run in list(paragraph.runs):
            paragraph._p.remove(run._r)
        run = paragraph.add_run(new)
        run.font.name = "Times New Roman"
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(12)
        found = True
        break

if not found:
    raise SystemExit("conclusion paragraph not found")

doc.core_properties.author = ""
doc.core_properties.last_modified_by = ""
doc.core_properties.comments = ""
doc.save(dst)
