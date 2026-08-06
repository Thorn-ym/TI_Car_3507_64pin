from __future__ import annotations

import copy
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor
from lxml import etree


BASE = Path(sys.argv[1]).resolve()
TEMPLATE = BASE / "inputs" / "template_converted.docx"
OUT = Path(sys.argv[2]).resolve()
MML_XSL = Path(r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL")

FONT_CN = "宋体"
FONT_HEADING = "黑体"
FONT_LATIN = "Times New Roman"


def set_run_font(run, name=FONT_CN, size=12, bold=None, italic=None):
    run.font.name = FONT_LATIN
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_para_spacing(p, line=22, before=0, after=0, first=0, keep=False):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line)
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.first_line_indent = Cm(first)
    pf.keep_with_next = keep


def clear_document(doc: Document):
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(22)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for style_name, font_name, size in (
        ("Heading 1", FONT_HEADING, 16),
        ("Heading 2", FONT_CN, 14),
        ("Heading 3", FONT_CN, 12),
    ):
        style = styles[style_name]
        style.font.name = FONT_LATIN
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        style.paragraph_format.line_spacing = Pt(22)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.keep_with_next = True

    if "Compact Table Text" not in [s.name for s in styles]:
        table_style = styles.add_style("Compact Table Text", WD_STYLE_TYPE.PARAGRAPH)
    else:
        table_style = styles["Compact Table Text"]
    table_style.font.name = FONT_LATIN
    table_style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_LATIN)
    table_style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    table_style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CN)
    table_style.font.size = Pt(9)
    table_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    table_style.paragraph_format.line_spacing = Pt(15)
    table_style.paragraph_format.space_before = Pt(0)
    table_style.paragraph_format.space_after = Pt(0)


def set_section_geometry(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.25)
    section.different_first_page_header_footer = False


def clear_story(story):
    for p in story.paragraphs:
        for run in list(p.runs):
            p._p.remove(run._r)
    for table in story.tables:
        table._element.getparent().remove(table._element)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_spacing(paragraph, line=14)
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=10.5)


def set_page_number_start(section, start=None):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    if start is not None:
        pg_num.set(qn("w:start"), str(start))


def add_body(doc, text, bold_prefix=None, first=0.74, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    set_para_spacing(p, first=first)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=12, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=12)
    else:
        r = p.add_run(text)
        set_run_font(r, size=12)
    return p


def add_title(doc, text, size=18, spacing_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, line=26, after=spacing_after, keep=True)
    r = p.add_run(text)
    set_run_font(r, FONT_HEADING, size=size, bold=True)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    r = p.add_run(text)
    set_run_font(r, FONT_HEADING, 16, True)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    r = p.add_run(text)
    set_run_font(r, FONT_CN, 14, True)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph(style="Heading 3")
    r = p.add_run(text)
    set_run_font(r, FONT_CN, 12, True)
    return p


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=80, bottom=70, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths_cm):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total_twips = int(sum(widths_cm) * 567)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 567)))
        grid.append(grid_col)
    for row in table.rows:
        for idx, (cell, width) in enumerate(zip(row.cells, widths_cm)):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 567)))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Cm(width)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, line=18, before=2, after=2, keep=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def add_table(doc, headers, rows, widths, font_size=9, header_fill="E7E6E6"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_widths(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.style = doc.styles["Compact Table Text"]
        r = p.add_run(text)
        set_run_font(r, FONT_CN, font_size, True)
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.style = doc.styles["Compact Table Text"]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_run_font(r, FONT_CN, font_size, False)
    set_para_spacing(doc.add_paragraph(), line=8)
    return table


def mathml_to_omml(mathml: str):
    transform = etree.XSLT(etree.parse(str(MML_XSL)))
    tree = etree.fromstring(mathml.encode("utf-8"))
    result = transform(tree)
    return parse_xml(etree.tostring(result.getroot()))


def add_equation(doc, mathml: str, number: str):
    p = doc.add_paragraph()
    set_para_spacing(p, line=24)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(8.0), WD_TAB_ALIGNMENT.CENTER)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(15.8), WD_TAB_ALIGNMENT.RIGHT)
    p.add_run("\t")
    run = p.add_run()
    run._r.append(mathml_to_omml(mathml))
    number_run = p.add_run(f"\t（{number}）")
    set_run_font(number_run, size=10.5)
    return p


def add_system_diagram(doc):
    table = doc.add_table(rows=5, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_widths(table, [5.0, 5.0, 5.0])
    texts = [
        ["USB摄像头\n动态端点与钢球图像", "树莓派5\n视觉测量/录像/图传", "外部接收显示存储装置"],
        ["↓", "USB CDC 50 Hz\n24字节+CRC", "↑ UDP/H.264预留"],
        ["七路红外+双编码器", "STM32F407\n滚球平衡控制", "M2006+C610\n摆杆执行机构"],
        ["↓", "BMI088加速度前馈\nCAN 1 Mbit/s", "↑ 编码器/速度/温度"],
        ["MSPM0G3507\n循迹、计时、显示", "四轮两驱差速底盘", "3S电池与分路电源"],
    ]
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=90, bottom=90)
            if r_idx in (0, 2, 4):
                set_cell_shading(cell, "F2F2F2")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.style = doc.styles["Compact Table Text"]
            r = p.add_run(texts[r_idx][c_idx])
            set_run_font(r, FONT_CN, 9.5, r_idx in (0, 2, 4))
    set_para_spacing(doc.add_paragraph(), line=8)


def add_placeholder(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_widths(table, [15.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8F8F8")
    set_cell_margins(cell, top=140, bottom=140, start=160, end=160)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, line=18)
    r = p.add_run(text)
    set_run_font(r, FONT_CN, 10.5, False)
    set_para_spacing(doc.add_paragraph(), line=8)


def add_reference(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, line=22, first=0)
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.first_line_indent = Cm(-0.74)
    r = p.add_run(text)
    set_run_font(r, size=12)


def build():
    doc = Document(TEMPLATE)
    clear_document(doc)
    configure_styles(doc)
    for section in doc.sections:
        set_section_geometry(section)
        clear_story(section.header)
        clear_story(section.footer)

    # 摘要页
    add_title(doc, "车载平衡滚球运动控制系统（H题）", 18, 8)
    add_title(doc, "摘  要", 16, 4)
    abstract = (
        "本设计面向H题车载平衡滚球运动控制系统，构建由七路红外循迹底盘、动态视觉测量、"
        "串级平衡控制和分路电源组成的车载系统。视觉端利用动态端点重建800×64归一化条带，"
        "融合空槽差分、暗体特征与Alpha-Beta跟踪，以50 Hz输出钢球位置；平衡端依据钢球滚动"
        "动力学，将位置PID、BMI088加速度前馈及M2006位置—速度双环结合；底盘端采用四轮两驱"
        "差速结构和七路加权质心循迹。视觉软件211项测试全部通过，完整链路约55～58帧/s，"
        "参考匹配素材空槽527帧零候选、带球975帧候选覆盖100%。理论单圈时间约16.7 s。系统已"
        "完成架构、接口及软件验证，钢球实测误差、停车偏差和整车圈时仍须以最终场地联调数据验收。"
    )
    add_body(doc, abstract, first=0.74)
    p = doc.add_paragraph()
    set_para_spacing(p, line=22, before=4)
    r1 = p.add_run("关键词：")
    set_run_font(r1, FONT_HEADING, 12, True)
    r2 = p.add_run("动态视觉；滚球平衡；串级控制；红外循迹；差速底盘")
    set_run_font(r2, FONT_CN, 12)

    # 正文从新页开始并重新编号
    body_sec = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_geometry(body_sec)
    body_sec.header.is_linked_to_previous = False
    body_sec.footer.is_linked_to_previous = False
    clear_story(body_sec.header)
    clear_story(body_sec.footer)
    add_page_field(body_sec.footer.paragraphs[0])
    set_page_number_start(body_sec, 1)

    add_title(doc, "车载平衡滚球运动控制系统（H题）", 18, 4)
    add_h1(doc, "一、系统方案")
    add_body(
        doc,
        "系统由循迹与底盘、视觉识别与图传、滚球平衡与摆杆执行、电源及人机交互组成。"
        "针对图像运算、毫秒级电机控制和10 ms循迹控制在资源与实时性上的差异，采用树莓派5、"
        "STM32F407和MSPM0G3507构成异构分布式控制架构。",
    )

    add_h2(doc, "1、主控制器件与电源方案的论证与选择")
    add_body(
        doc,
        "方案一采用单片机集中完成视觉、滚球和底盘控制，硬件简单但难以承担60帧/s图像处理；"
        "方案二采用单板计算机统一控制，图像能力充足但硬实时、PWM与故障隔离能力不足；方案三"
        "按任务时间尺度分配控制器：树莓派5处理OpenCV视觉及显示存储，STM32F407以1 kHz完成"
        "滚球与M2006串级控制，MSPM0G3507以100 Hz完成七路循迹、双编码器和计时显示。故选方案三。",
    )
    add_body(
        doc,
        "电源方面，线性稳压损耗大，单路开关电源易使电机纹波耦合至相机和传感器；本设计采用"
        "3S锂电池形成约11.1 V动力母线，C610和底盘H桥由动力支路供电，独立降压形成5 V视觉/"
        "图传电源及3.3 V逻辑电源，各支路共地并配置保险、反接保护、储能电容与高频去耦。9.5 V"
        "持续300 ms触发欠压停机。该方案兼顾M2006现有12 V运行状态、树莓派瞬态负载和传感器低噪声要求。",
    )

    add_h2(doc, "2、视觉识别与图传方案的论证与选择")
    add_body(
        doc,
        "固定ROI方案计算量小，但摆杆转动或车体振动后会发生区域截断；YOLO等深度学习方案适应性强，"
        "但需大量标注样本，树莓派CPU部署开销较大，且检测框仍需轴向标定。最终采用动态端点标记、"
        "透视归一化和一维融合检测：左端单标记、右端双标记实时重建摆杆坐标系，将图像映射为"
        "800×64条带，再融合三姿态空槽参考差分、梯度和局部暗体响应。USB摄像头工作于"
        "640×480、MJPG、60帧/s；树莓派5兼顾识别、录像和网络预览。图传预留H.264/UDP 25帧/s、"
        "4 Mbit/s、GOP 30、端口5600和队列深度2，外部计算机负责显示、存储及回放。",
    )

    add_h2(doc, "3、滚球平衡与摆杆执行方案的论证与选择")
    add_body(
        doc,
        "单闭环位置PID实现简单，但小车加减速时只能在钢球位移后纠偏；LQR能综合多状态，然而对"
        "非对称连杆、摩擦和完整状态估计依赖较强。最终选择钢球位置PID、滚动动力学反算、BMI088"
        "纵向加速度前馈以及电机位置—速度双环的复合方案，使视觉反馈消除稳态偏差，前馈提前抵消"
        "底盘扰动，快慢环分离便于现场整定。执行机构比较舵机、有刷减速电机和M2006+C610后，选用"
        "36:1 M2006 P36无刷减速电机及C610电调；其8192计数编码器、CAN反馈和温度信息适合高频微调，"
        "连续转矩1 N·m，正常输出轴速度限制为±45 r/min。",
    )

    add_h2(doc, "4、循迹检测与底盘方案的论证与选择")
    add_body(
        doc,
        "线阵CCD能够提供连续灰度与曲率信息，但处理和曝光链路复杂，且题目明确规定循迹只能使用"
        "红外光电模块。故选七路反射式红外数字阵列，S1～S7横向布置，中间S4对准车体中心；多通道"
        "状态既可估计横向偏差，也可识别A点宽启停线。底盘采用四轮两驱差速结构，两只主动轮独立"
        "编码调速，两只万向轮承载随动，机械链短、四点支承稳定，适合两段1.5 m直线和半径0.5 m"
        "半圆弧组成的环形路线。驱动选用双路TB6612FNG类H桥，终点采用短路制动。",
    )
    add_caption(doc, "表1  系统各子系统最终选型")
    add_table(
        doc,
        ["子系统", "比较方案", "最终方案", "选型结论"],
        [
            ["控制器", "单MCU/单板机/异构分布", "树莓派5+STM32F407+MSPM0G3507", "算力与实时控制分离"],
            ["视觉", "固定ROI/深度学习/动态几何", "动态标记+归一化条带", "实时、可解释、适应摆杆运动"],
            ["平衡", "单PID/LQR/串级复合", "PID+逆动力学+IMU前馈", "抗扰与调试效率兼顾"],
            ["摆杆驱动", "舵机/有刷H桥/M2006+C610", "M2006 P36+C610", "反馈完整、CAN抗干扰"],
            ["循迹底盘", "CCD/红外；舵向/差速", "七路红外+四轮两驱", "符合题目限制、支承稳定"],
            ["电源", "线性/单路DC-DC/分路DC-DC", "3S母线+动力/5 V/3.3 V分路", "降低动力噪声耦合"],
        ],
        [2.2, 4.1, 4.2, 5.5],
        8.5,
    )

    add_h1(doc, "二、系统理论分析与计算")
    add_h2(doc, "1、视觉测量原理与计算")
    add_body(
        doc,
        "设左右端点为P_L、P_R，右端双标记中心取两候选点均值；由端点差建立轴向单位向量u和"
        "法向量n，并用四点单应矩阵H把动态四边形映射为固定条带。端点距离限制430～610像素，"
        "摆杆角度限制±12°，帧间跳变不超过35像素；标记白阈值150，面积4～80 px²、圆度不低于"
        "0.38、长宽比不大于1.8、黑环比例不低于0.30，右端点距8～30像素、方向90°±35°。",
    )
    add_equation(
        doc,
        '<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><mi>u</mi><mo>=</mo><mfrac><mrow><msub><mi>P</mi><mi>R</mi></msub><mo>−</mo><msub><mi>P</mi><mi>L</mi></msub></mrow><mrow><mo>∥</mo><msub><mi>P</mi><mi>R</mi></msub><mo>−</mo><msub><mi>P</mi><mi>L</mi></msub><mo>∥</mo></mrow></mfrac><mo>,</mo><mspace width="0.5em"/><mi>n</mi><mo>=</mo><mo>(</mo><mo>−</mo><msub><mi>u</mi><mi>y</mi></msub><mo>,</mo><msub><mi>u</mi><mi>x</mi></msub><mo>)</mo></mrow></math>',
        "2-1",
    )
    add_equation(
        doc,
        '<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><mi>λ</mi><msup><mi>p</mi><mo>′</mo></msup><mo>=</mo><mi>H</mi><mi>p</mi></mrow></math>',
        "2-2",
    )
    add_body(
        doc,
        "钢球检测将灰度差分与梯度差分按1.0:0.65融合，9像素平滑后筛选峰值18、峰宽8～80像素"
        "和跟踪半径120像素；局部暗度路径采用101像素基线核、阈值8.0和3～60像素峰宽作为参考"
        "失配时的补充。低、中、高参考摆杆长度为501.88、558.79、589.59像素。钢球中心投影到"
        "摆杆轴后，以t_left=0.9495781064、t_right=0.0514231734换算24.0 cm有效行程，理论采样"
        "尺度约0.30 mm/像素。",
    )
    add_equation(
        doc,
        '<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><mi>t</mi><mo>=</mo><mfrac><mrow><mo>(</mo><msub><mi>P</mi><mi>B</mi></msub><mo>−</mo><msub><mi>P</mi><mi>L</mi></msub><mo>)</mo><mo>·</mo><mo>(</mo><msub><mi>P</mi><mi>R</mi></msub><mo>−</mo><msub><mi>P</mi><mi>L</mi></msub><mo>)</mo></mrow><msup><mrow><mo>∥</mo><msub><mi>P</mi><mi>R</mi></msub><mo>−</mo><msub><mi>P</mi><mi>L</mi></msub><mo>∥</mo></mrow><mn>2</mn></msup></mfrac><mo>,</mo><mspace width="0.5em"/><mi>x</mi><mo>=</mo><msub><mi>L</mi><mi>eff</mi></msub><mfrac><mrow><mi>t</mi><mo>−</mo><msub><mi>t</mi><mi>left</mi></msub></mrow><mrow><msub><mi>t</mi><mi>right</mi></msub><mo>−</mo><msub><mi>t</mi><mi>left</mi></msub></mrow></mfrac></mrow></math>',
        "2-3",
    )
    add_body(
        doc,
        "Alpha-Beta跟踪器取α=0.75、β=0.12，依据恒速度模型校正位置和速度；漏检不超过3帧时"
        "输出PREDICTED，超限即LOST。该处理提高连续性，但预测状态不能等同直接检测。",
    )
    add_equation(
        doc,
        '<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><msup><mover><mi>x</mi><mo>^</mo></mover><mo>−</mo></msup><msub><mrow></mrow><mi>k</mi></msub><mo>=</mo><msub><mover><mi>x</mi><mo>^</mo></mover><mrow><mi>k</mi><mo>−</mo><mn>1</mn></mrow></msub><mo>+</mo><msub><mover><mi>v</mi><mo>^</mo></mover><mrow><mi>k</mi><mo>−</mo><mn>1</mn></mrow></msub><mi>Δt</mi><mo>,</mo><mspace width="0.5em"/><msub><mover><mi>x</mi><mo>^</mo></mover><mi>k</mi></msub><mo>=</mo><msup><mover><mi>x</mi><mo>^</mo></mover><mo>−</mo></msup><msub><mrow></mrow><mi>k</mi></msub><mo>+</mo><mi>α</mi><msub><mi>r</mi><mi>k</mi></msub></mrow></math>',
        "2-4",
    )

    add_h2(doc, "2、滚球平衡与电机串级控制计算")
    add_body(
        doc,
        "钢球近似为均匀实心球，在凹槽内无滑动滚动。设钢球轴向位置为x、摆杆角为θ、小车纵向"
        "加速度为a_c，则平动与转动惯量共同形成5/7等效系数。",
    )
    add_equation(
        doc,
        '<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><mover><mi>x</mi><mo>¨</mo></mover><mo>=</mo><mo>−</mo><mfrac><mn>5</mn><mn>7</mn></mfrac><mo>(</mo><mi>g</mi><mi>sin</mi><mi>θ</mi><mo>+</mo><msub><mi>a</mi><mi>c</mi></msub><mi>cos</mi><mi>θ</mi><mo>)</mo></mrow></math>',
        "2-5",
    )
    add_body(
        doc,
        "位置外环按目标位置与视觉位置之差产生目标加速度，参数K_px=4.0、K_ix=0.35、K_dx=2.2，"
        "积分限幅0.35 m/s²、输出限幅1.5 m/s²，速度低通12 Hz、微分低通10 Hz。忽略小积分项时，"
        "自然频率2 rad/s、阻尼比0.55，理论2%调节时间约3.64 s、峰值超调约12.6%，与题目5 s"
        "往返要求具有理论一致性。",
    )
    add_equation(
        doc,
        '<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><msub><mi>a</mi><mi>d</mi></msub><mo>=</mo><msub><mi>K</mi><mi>px</mi></msub><msub><mi>e</mi><mi>x</mi></msub><mo>+</mo><msub><mi>K</mi><mi>ix</mi></msub><mo>∫</mo><msub><mi>e</mi><mi>x</mi></msub><mi>dt</mi><mo>−</mo><msub><mi>K</mi><mi>dx</mi></msub><msub><mover><mi>x</mi><mo>˙</mo></mover><mi>f</mi></msub></mrow></math>',
        "2-6",
    )
    add_body(
        doc,
        "由动力学逆解得到摆杆目标角，并叠加BMI088纵向加速度前馈。目标角限幅±4.5°，静止平台"
        "最大钢球加速度约0.550 m/s²，可完整抵消的纵向加速度约±0.772 m/s²。BMI088采用8 Hz"
        "低通，时间常数约19.9 ms，上电1000点零偏标定约1 s。",
    )
    add_equation(
        doc,
        '<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><msub><mi>θ</mi><mi>d</mi></msub><mo>=</mo><mi>arcsin</mi><mo>(</mo><mfrac><mrow><mo>−</mo><msub><mi>a</mi><mi>d</mi></msub></mrow><mrow><mfrac><mn>5</mn><mn>7</mn></mfrac><msqrt><mrow><msup><mi>g</mi><mn>2</mn></msup><mo>+</mo><msup><msub><mi>a</mi><mi>f</mi></msub><mn>2</mn></msup></mrow></msqrt></mrow></mfrac><mo>)</mo><mo>−</mo><mi>arctan</mi><mo>(</mo><mfrac><msub><mi>a</mi><mi>f</mi></msub><mi>g</mi></mfrac><mo>)</mo></mrow></math>',
        "2-7",
    )
    add_body(
        doc,
        "M2006转子编码器0～8191计数，经36:1减速后每计数对应输出轴0.0012207°。实测水平累计"
        "值4552，右端最低13760（相对+11.240234°），右端最高2145（相对−2.938232°）；采用分段"
        "线性插值描述非对称连杆，±4.5°软限位对应电机约+10.1162°至−2.6444°。位置PD参数"
        "7.0/0/0.12，速度限幅±45 r/min；速度PI参数1.2/8.0/0，积分限幅1200、电流指令限幅"
        "±2500、最小有效值160，积分时间常数0.15 s。",
    )

    add_h2(doc, "3、循迹、差速运动学与圈时计算")
    add_body(
        doc,
        "七路状态b_i对应权值+50、+33、+16、0、−16、−33、−50。有效通道数为n时，按加权质心"
        "计算路径偏差；外环PID取K_p=0.9、K_i=0.05、K_d=0.2，并以0.04倍MPU6050 Z轴角速度"
        "形成偏航阻尼，修正量限幅±30计数。左右速度内环参数分别为180/0.35/26和180/0.28/18，"
        "PWM限幅±3199，频率约10 kHz。",
    )
    add_equation(
        doc,
        '<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><mi>e</mi><mo>=</mo><mfrac><mrow><mo>∑</mo><msub><mi>b</mi><mi>i</mi></msub><msub><mi>w</mi><mi>i</mi></msub></mrow><mi>n</mi></mfrac><mo>,</mo><mspace width="0.5em"/><mi>v</mi><mo>=</mo><mfrac><mrow><msub><mi>v</mi><mi>R</mi></msub><mo>+</mo><msub><mi>v</mi><mi>L</mi></msub></mrow><mn>2</mn></mfrac><mo>,</mo><mspace width="0.5em"/><mi>ω</mi><mo>=</mo><mfrac><mrow><msub><mi>v</mi><mi>R</mi></msub><mo>−</mo><msub><mi>v</mi><mi>L</mi></msub></mrow><mi>B</mi></mfrac></mrow></math>',
        "2-8",
    )
    add_body(
        doc,
        "轨道中心线长度为3+π≈6.142 m。五次整圈编码器中值为左49712、右39741计数，巡航目标"
        "31计数/10 ms对应名义0.431 m/s；采用1.30 s平滑加速、圈进度86%～97%减速及末段"
        "11计数/10 ms（约0.153 m/s）的速度规划，离散推算单圈约16.7 s，理论上低于20 s要求。"
        "A点离开确认5周期、最短运行100周期、回到启停线连续2周期确认，抑制起点误判。",
    )

    add_h1(doc, "三、电路与程序设计")
    add_h2(doc, "1、系统总体框图与整体工作流程")
    add_body(
        doc,
        "按键启动后，循迹控制器开始10 ms计时并驱动底盘；视觉端持续采集摆杆与钢球图像，形成"
        "带状态和时间戳的位置帧；平衡控制器仅在CRC、序号、时效和状态有效时更新钢球外环，同时"
        "读取BMI088和M2006反馈，经角度反算及双环控制向C610下发电流。回到A点后底盘短刹并停止"
        "计时；任何关键故障均撤销相应驱动输出。系统总体框图如图1所示。",
    )
    add_system_diagram(doc)
    add_caption(doc, "图1  系统总体框图")

    add_h2(doc, "2、子系统硬件电路设计")
    add_h3(doc, "（1）视觉识别与图传电路")
    add_body(
        doc,
        "USB摄像头接树莓派5的USB主机口，固定曝光值157、亮度4并关闭动态帧率，以保持参考差分"
        "条件稳定。树莓派通过USB Device CDC与STM32F407连接，辅助显示、录像和图传采用有界队列，"
        "避免网络或存储阻塞识别线程。摄像头和图传发送端牢固安装于车体，画面覆盖完整25 cm摆杆；"
        "接收计算机置于线路外完成实时显示、存储和回放。",
    )
    add_h3(doc, "（2）滚球平衡与电机驱动电路")
    add_body(
        doc,
        "STM32F407以SPI1约2.625 MHz读取BMI088，以CAN1 1 Mbit/s连接C610，电机ID为1。控制命令"
        "标准帧ID为0x200，反馈ID为0x201；反馈包含编码器、转速、电流和温度。ADC3以100 Hz采集"
        "电池电压。CAN端配置匹配的收发器、终端和可靠共地，电机动力线与USB、SPI及传感器线分离。"
        "摆杆左端铰接高度不低于5 cm，右端由连杆机构驱动，并设置机械实体挡块。",
    )
    add_h3(doc, "（3）循迹与底盘电路")
    add_body(
        doc,
        "MSPM0G3507系统时钟32 MHz，S1～S7接PB12～PB18并内部上拉；左编码器由TIMG8硬件QEI"
        "采集，右编码器由PB10/PB11双边沿中断解码。TIMA0两通道输出约10 kHz PWM至TB6612FNG，"
        "STBY独立控制；MPU6050经I2C1 100 kHz连接，SSD1306 128×64 OLED采用GPIO开漏模拟I2C，"
        "显示时间、圈次和里程。所有MCU输入均限制在3.3 V以内。",
    )
    add_h3(doc, "（4）电源与抗干扰电路")
    add_body(
        doc,
        "3S电池经总开关、保险与反接保护后分为动力和逻辑支路；M2006/C610与底盘电机端就近配置"
        "储能及高频去耦，树莓派和摄像头采用独立5 V大电流降压，MCU与传感器采用低纹波5 V/3.3 V"
        "支路。各支路单点汇接，共模地线短而粗；USB、编码器和I2C线远离H桥开关节点。",
    )

    add_h2(doc, "3、程序设计、通信协议与异常保护")
    add_body(
        doc,
        "视觉程序依次执行端点检测、几何校验、条带重建、候选融合、位置标定、Alpha-Beta跟踪及"
        "50 Hz发送；滚球程序执行视觉帧解析、模式状态机、位置PID、前馈反算、三点标定、位置环、"
        "速度环与安全判定；底盘程序每10 ms完成编码器增量、七路采样、速度规划、循迹PID、双轮"
        "速度PID及启停线状态机。非关键OLED、GUI、Web和录像均采用降级或限帧策略。",
    )
    add_caption(doc, "表2  关键接口与协议")
    add_table(
        doc,
        ["接口", "配置", "数据内容", "有效性与失效处理"],
        [
            ["视觉→平衡", "USB CDC，115200-8N1显示参数，50 Hz", "24字节小端；头A5 5A、版本1、序号、时间戳、位置0.1 mm/LSB、速度、置信度、状态、CRC-16", "预测最多100 ms且权重50%；150 ms超时；无效位置−32768"],
            ["平衡→C610", "CAN1，1 Mbit/s，ID 0x200/0x201", "电流命令；编码器、转速、反馈电流、温度", "反馈100 ms超时停机；80 ℃过温锁存"],
            ["循迹采集", "GPIO/QEI，100 Hz", "7位黑线状态、双轮增量、启停按键", "脱线1周期撤销PWM；终点连续2周期确认"],
            ["图传", "H.264/UDP，25帧/s，4 Mbit/s", "覆盖摆杆的实时视频，端口5600", "队列2优先低延迟；断网不阻塞视觉主链路"],
        ],
        [2.3, 3.5, 5.7, 4.5],
        8.2,
    )
    add_caption(doc, "表3  关键参数配置")
    add_table(
        doc,
        ["模块", "采样/范围", "控制或检测参数", "保护与资源"],
        [
            ["视觉", "640×480@60；条带800×64；有效行程0～240 mm", "差分权重1/0.65；峰阈值18；峰宽8～80；α/β=0.75/0.12；串口50 Hz", "左右端保持2/3帧；球漏检3帧；写超时0.02 s；重连1 s；录像/图传队列8/2"],
            ["平衡", "控制1 kHz；角度±4.5°；速度±45 r/min", "球PID 4.0/0.35/2.2；位置PD 7/0/0.12；速度PI 1.2/8；电流±2500", "视觉150 ms；CAN100 ms；IMU30 ms；9.5 V/300 ms；80 ℃；Flash 40.8 kB、RAM 10.2 kB"],
            ["底盘", "控制100 Hz；PWM约10 kHz；计时0.01 s", "循迹0.9/0.05/0.2；左右速度180/0.35/26、180/0.28/18；巡航31计数", "脱线停机；STBY关断；I2C/OLED自动恢复；Flash 22344 B、SRAM 4507 B"],
        ],
        [2.0, 3.8, 6.3, 3.9],
        8.0,
    )
    add_body(
        doc,
        "异常保护包括：视觉CRC、长度、状态、范围、序号和时效校验；摆杆±4.5°软限位及电机"
        "−2.938232°/+11.240234°实测端点锁存；零点再次捕获窗口0.588°；电流、速度、积分和"
        "加速度多级限幅；模式切换清零控制器；底盘脱线、编码器停滞和双轮失配停机。BMI088失效时"
        "前馈置零但保留视觉闭环，OLED或网络失效不影响主控制。现阶段尚未实现独立堵转持续时间"
        "判据，最终实物应增加机械保险或熔断保护。",
    )
    add_placeholder(
        doc,
        "程序流程图占位：图2 主程序流程图；图3 视觉识别流程图；图4 滚球平衡控制流程图；"
        "图5 循迹与停车流程图。正式提交时可在保持图号与版式不变的条件下替换为最终流程图。",
    )

    add_h1(doc, "四、测试方案与测试结果")
    add_h2(doc, "1、测试方案")
    add_body(
        doc,
        "硬件测试依次检查分路电源纹波、相机60帧/s采集、七路红外极性与阈值、编码器方向、"
        "TB6612双轮PWM、M2006正负小角度响应、机械挡块、CAN反馈、欠压及过温停机；软件测试覆盖"
        "视觉几何与候选、CRC和粘包恢复、跟踪状态、PID限幅与抗饱和、三点标定、模式切换、启停线"
        "状态机和外设降级。系统联调在标准1.8±0.2 cm黑线场地进行，分别完成静态±5 cm往返、"
        "AB段、中心整圈、任意位置整圈和图传显示存储测试。",
    )
    add_body(
        doc,
        "静态滚球至少重复5次，以钢球进入目标±1 cm并保持为调节完成，记录总时间、最大误差和"
        "平均误差；循迹及停车至少重复10次，记录单圈时间、AB时间、停车偏差、脱线次数和成功率；"
        "整车测试同步保存视频与Ozone/串口遥测，统计视觉数据年龄、钢球最大误差、摆杆限幅次数、"
        "电流峰值及安全停机次数。主要仪器包括示波器、数字万用表、倾角仪、标尺、外部计时录像"
        "与显示存储计算机。",
    )

    add_h2(doc, "2、测试结果及分析")
    add_caption(doc, "表4  已有测试结果与题目指标对照")
    add_table(
        doc,
        ["测试项目", "题目/设计指标", "现有结果", "结论"],
        [
            ["视觉软件", "功能与异常处理正确", "211项全部通过，用时2.59 s", "已验证"],
            ["相机采集", "640×480、60帧/s稳定", "1807帧/30.10 s；60.02帧/s；失败0", "满足"],
            ["完整视觉链路", "单帧低于16.67 ms", "55～58帧/s；P95 7.1～7.5 ms；P99 7.5～8.5 ms", "满足实时性"],
            ["离线识别", "空槽低误检、带球连续", "参考匹配空槽527帧零候选；带球975帧候选覆盖100%；端点100%", "样本范围满足；非严格精度"],
            ["广泛动态素材", "适应姿态变化", "12组端点100%；全范围压力素材候选覆盖95.33%", "方案有效"],
            ["平衡固件", "1 kHz、50 Hz输入、1 Mbit/s CAN", "全量交叉编译；Flash约40.8 kB、RAM约10.2 kB", "软件与资源满足"],
            ["底盘固件", "100 Hz闭环、计时显示", "构建通过；Flash 22344 B、SRAM 4507 B", "软件与资源满足"],
            ["理论圈时", "单圈≤20 s", "速度规划推算约16.7 s", "有理论余量，非实测"],
            ["静态±5 cm", "≤5 s，最大误差≤1 cm", "源材料未提供规范化实测记录", "待整机验收"],
            ["AB/整圈球位", "AB≤8 s；整圈≤30 s；误差≤1 cm", "源材料未提供整车实测数据", "待整机验收"],
            ["停车与图传", "停车≤2 cm；稳定显示、记录、回放", "停车偏差及无线图传端到端数据未提供", "待场地/实机验收"],
        ],
        [3.0, 3.5, 5.6, 3.9],
        8.0,
    )
    add_body(
        doc,
        "测试表明视觉算法、嵌入式实时性、通信协议、控制软件和资源占用达到当前设计阶段要求。"
        "空槽零候选与带球候选覆盖证明检测链路具有稳定候选输出，但未建立逐帧物理真值，不能将"
        "100%候选覆盖解释为100%识别准确率；0.30 mm/像素仅是理论采样尺度，也不能替代±1 cm"
        "位置误差验收。底盘16.7 s为依据编码器标定和速度曲线的计算值，停车偏差、实际圈时及"
        "满载弯道成功率仍应以标准场地数据为准。",
    )
    add_body(
        doc,
        "综合现有证据，本设计已完成H题所需的系统架构、红外循迹、视觉测量、滚球控制、计时显示"
        "和安全保护设计，模块软件与实时性指标达到预期，并具备完成赛题的理论条件；但由于缺少"
        "静态滚球误差、AB段时间、整圈钢球误差、停车偏差及无线图传实机数据，现阶段不能正式判定"
        "整机已全部达到题目要求。补齐表4所列实测项目且最差值满足限值后，方可作出“本设计达到"
        "H题全部要求”的最终验收结论。",
    )

    # 参考文献与附录另起页，页码连续
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_h1(doc, "五、参考文献")
    refs = [
        "[1] TEXAS INSTRUMENTS. MSPM0G350x Mixed-Signal Microcontrollers Datasheet[EB/OL]. Dallas: Texas Instruments, 2025.",
        "[2] STMICROELECTRONICS. STM32F405xx/STM32F407xx Datasheet[EB/OL]. Geneva: STMicroelectronics, 2025.",
        "[3] DJI. RoboMaster M2006 P36直流无刷减速电机与C610电调使用说明[Z]. 深圳: 深圳市大疆创新科技有限公司.",
        "[4] BRADSKI G, KAEHLER A. Learning OpenCV: Computer Vision with the OpenCV Library[M]. Sebastopol: O'Reilly Media, 2008.",
        "[5] FRANKLIN G F, POWELL J D, EMAMI-NAEINI A. Feedback Control of Dynamic Systems[M]. 8th ed. Boston: Pearson, 2019.",
        "[6] CORKE P. Robotics, Vision and Control: Fundamental Algorithms in MATLAB[M]. 2nd ed. Cham: Springer, 2017.",
    ]
    for ref in refs:
        add_reference(doc, ref)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_h1(doc, "附录1：电路原理图")
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_h1(doc, "附录2：源程序")

    # 全部节应用页面规范，清空页眉；正文之后页码保持连续。
    for idx, section in enumerate(doc.sections):
        set_section_geometry(section)
        clear_story(section.header)
        if idx == 0:
            clear_story(section.footer)
        elif not section.footer.paragraphs[0].text and not section.footer.paragraphs[0]._p.xpath('.//w:fldChar'):
            add_page_field(section.footer.paragraphs[0])

    doc.core_properties.title = "车载平衡滚球运动控制系统（H题）设计报告"
    doc.core_properties.subject = "全国大学生电子设计竞赛广东赛区设计报告"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.comments = ""
    doc.settings.element.get_or_add_updateFields().set(qn("w:val"), "true")
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build()
