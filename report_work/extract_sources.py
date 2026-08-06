from __future__ import annotations

import json
import sys
from pathlib import Path

import pdfplumber
from docx import Document
from docx.oxml.ns import qn


def twips_to_cm(value: int | None) -> float | None:
    if value is None:
        return None
    return round(value / 567.0, 3)


def font_name(run) -> str | None:
    rpr = run._element.rPr
    if rpr is None or rpr.rFonts is None:
        return run.font.name
    return (
        rpr.rFonts.get(qn("w:eastAsia"))
        or rpr.rFonts.get(qn("w:ascii"))
        or run.font.name
    )


def para_record(paragraph) -> dict:
    fmt = paragraph.paragraph_format
    runs = []
    for run in paragraph.runs:
        if not run.text:
            continue
        runs.append(
            {
                "text": run.text,
                "font": font_name(run),
                "size_pt": round(run.font.size.pt, 2) if run.font.size else None,
                "bold": run.bold,
                "italic": run.italic,
                "underline": bool(run.underline) if run.underline is not None else None,
            }
        )
    return {
        "text": paragraph.text,
        "style": paragraph.style.name if paragraph.style else None,
        "alignment": str(paragraph.alignment) if paragraph.alignment is not None else None,
        "line_spacing": str(fmt.line_spacing) if fmt.line_spacing is not None else None,
        "space_before_pt": round(fmt.space_before.pt, 2) if fmt.space_before else None,
        "space_after_pt": round(fmt.space_after.pt, 2) if fmt.space_after else None,
        "first_line_indent_cm": twips_to_cm(fmt.first_line_indent.twips) if fmt.first_line_indent else None,
        "left_indent_cm": twips_to_cm(fmt.left_indent.twips) if fmt.left_indent else None,
        "keep_with_next": fmt.keep_with_next,
        "page_break_before": fmt.page_break_before,
        "runs": runs,
    }


def extract_docx(path: Path) -> dict:
    doc = Document(path)
    sections = []
    for section in doc.sections:
        sections.append(
            {
                "page_width_cm": twips_to_cm(section.page_width.twips),
                "page_height_cm": twips_to_cm(section.page_height.twips),
                "top_margin_cm": twips_to_cm(section.top_margin.twips),
                "bottom_margin_cm": twips_to_cm(section.bottom_margin.twips),
                "left_margin_cm": twips_to_cm(section.left_margin.twips),
                "right_margin_cm": twips_to_cm(section.right_margin.twips),
                "header_distance_cm": twips_to_cm(section.header_distance.twips),
                "footer_distance_cm": twips_to_cm(section.footer_distance.twips),
                "different_first_page_header_footer": section.different_first_page_header_footer,
                "header": [para_record(p) for p in section.header.paragraphs],
                "footer": [para_record(p) for p in section.footer.paragraphs],
            }
        )

    tables = []
    for index, table in enumerate(doc.tables, 1):
        rows = []
        for row in table.rows:
            rows.append([[para_record(p) for p in cell.paragraphs] for cell in row.cells])
        tables.append({"index": index, "style": table.style.name if table.style else None, "rows": rows})

    style_records = []
    for style in doc.styles:
        if style.type != 1:
            continue
        pf = style.paragraph_format
        style_records.append(
            {
                "name": style.name,
                "base_style": style.base_style.name if style.base_style else None,
                "font": style.font.name,
                "size_pt": round(style.font.size.pt, 2) if style.font.size else None,
                "bold": style.font.bold,
                "alignment": str(pf.alignment) if pf.alignment is not None else None,
                "line_spacing": str(pf.line_spacing) if pf.line_spacing is not None else None,
                "space_before_pt": round(pf.space_before.pt, 2) if pf.space_before else None,
                "space_after_pt": round(pf.space_after.pt, 2) if pf.space_after else None,
                "first_line_indent_cm": twips_to_cm(pf.first_line_indent.twips) if pf.first_line_indent else None,
            }
        )

    return {
        "path": str(path),
        "sections": sections,
        "paragraphs": [para_record(p) for p in doc.paragraphs],
        "tables": tables,
        "styles": style_records,
        "inline_shapes": len(doc.inline_shapes),
    }


def extract_pdf(path: Path) -> dict:
    pages = []
    with pdfplumber.open(path) as pdf:
        for index, page in enumerate(pdf.pages, 1):
            pages.append(
                {
                    "page": index,
                    "width": page.width,
                    "height": page.height,
                    "text": page.extract_text(x_tolerance=2, y_tolerance=3) or "",
                    "tables": page.extract_tables(),
                }
            )
    return {"path": str(path), "pages": pages}


def main() -> None:
    base = Path(sys.argv[1]).resolve()
    out = Path(sys.argv[2]).resolve()
    out.mkdir(parents=True, exist_ok=True)

    for name in ("template_converted.docx", "vision.docx", "chassis.docx"):
        data = extract_docx(base / name)
        (out / f"{Path(name).stem}.json").write_text(
            json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8"
        )

    pdf_data = extract_pdf(base / "problem.pdf")
    (out / "problem.json").write_text(
        json.dumps(pdf_data, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    balance = (base / "balance_motor.txt").read_text(encoding="utf-8-sig")
    (out / "balance_motor.txt").write_text(balance, encoding="utf-8")


if __name__ == "__main__":
    main()
