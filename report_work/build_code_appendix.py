from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output"
OUTPUT_PATH = OUTPUT_DIR / "设计报告核心代码摘录.docx"


def read_range(filename: str, start: int, end: int) -> list[str]:
    lines = (ROOT / filename).read_text(encoding="utf-8").splitlines()
    return lines[start - 1 : end]


def set_east_asia_font(run, east_asia: str) -> None:
    run._element.get_or_add_rPr().get_or_add_rFonts().set(
        qn("w:eastAsia"), east_asia
    )


def set_style_font(style, ascii_font: str, east_asia: str, size: Pt) -> None:
    style.font.name = ascii_font
    style.font.size = size
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:cs"), ascii_font)



def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    set_east_asia_font(run, "宋体")

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.extend([begin, instruction, separate, text, end])


def wrap_selected_lines(name: str, lines: list[str]) -> list[str]:
    result: list[str] = []
    for line in lines:
        stripped = line.strip()

        if name == "line_tracker.c":
            if stripped.startswith("uint8_t is_set ="):
                result.extend(
                    [
                        "    uint8_t is_set =",
                        "        (DL_GPIO_readPins(line_ports[index],",
                        "                          line_pins[index]) != 0U) ?",
                        "        1U : 0U;",
                    ]
                )
                continue
            if stripped.startswith("is_active ="):
                result.extend(
                    [
                        "    is_active = (g_line.active_low != 0U) ?",
                        "        (uint8_t)(is_set == 0U) : is_set;",
                    ]
                )
                continue
            if stripped.startswith("uint8_t left_right_angle ="):
                result.extend(
                    [
                        "      uint8_t left_right_angle =",
                        "          ((active_mask & 0x0FU) == 0x0FU) ? 1U : 0U;",
                    ]
                )
                continue
            if stripped.startswith("uint8_t right_right_angle ="):
                result.extend(
                    [
                        "      uint8_t right_right_angle =",
                        "          ((active_mask & 0x78U) == 0x78U) ? 1U : 0U;",
                    ]
                )
                continue
            if stripped == "else if ((right_right_angle != 0U) && (left_right_angle == 0U))":
                result.extend(
                    [
                        "      else if ((right_right_angle != 0U) &&",
                        "               (left_right_angle == 0U))",
                    ]
                )
                continue

        if name == "car_control.c":
            if stripped.startswith("float error = (float)motor->target_counts"):
                result.extend(
                    [
                        "  float error = (float)motor->target_counts -",
                        "                (float)motor->measured_counts;",
                    ]
                )
                continue
            if stripped.startswith("motor->pid.integral = Car_LimitFloat"):
                result.extend(
                    [
                        "  motor->pid.integral = Car_LimitFloat(",
                        "      motor->pid.integral, motor->pid.integral_limit);",
                    ]
                )
                continue
            if stripped.startswith("pid->integral = Car_LimitFloat"):
                result.extend(
                    [
                        "  pid->integral = Car_LimitFloat(",
                        "      pid->integral, pid->integral_limit);",
                    ]
                )
                continue

        if name == "competition_tasks.c":
            if stripped.startswith("OdometerControlProgress_t progress ="):
                result.extend(
                    [
                        "    OdometerControlProgress_t progress =",
                        "        Odometer_GetControlProgress();",
                    ]
                )
                continue
            if stripped.startswith("uint32_t blend = CompetitionTasks_SmoothStep"):
                result.extend(
                    [
                        "      uint32_t blend = CompetitionTasks_SmoothStep(",
                        "          position, span);",
                    ]
                )
                continue

        result.append(line.rstrip())

    return result


def create_document() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)

    normal = doc.styles["Normal"]
    set_style_font(normal, "SimSun", "宋体", Pt(12))
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    code_title = doc.styles.add_style("Code Caption Academic", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(code_title, "SimHei", "黑体", Pt(12))
    code_title.font.bold = True
    code_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    code_title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    code_title.paragraph_format.space_before = Pt(0)
    code_title.paragraph_format.space_after = Pt(5)
    code_title.paragraph_format.keep_with_next = True

    code_style = doc.styles.add_style("Academic Code", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(code_style, "Consolas", "Consolas", Pt(10.5))
    code_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    code_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    code_style.paragraph_format.left_indent = Cm(0.74)
    code_style.paragraph_format.first_line_indent = Cm(0)
    code_style.paragraph_format.space_before = Pt(0)
    code_style.paragraph_format.space_after = Pt(0)
    code_style.paragraph_format.widow_control = False

    explanation = doc.styles.add_style(
        "Code Explanation Academic", WD_STYLE_TYPE.PARAGRAPH
    )
    set_style_font(explanation, "SimSun", "宋体", Pt(9))
    explanation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    explanation.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    explanation.paragraph_format.first_line_indent = Cm(0.63)
    explanation.paragraph_format.space_before = Pt(5)
    explanation.paragraph_format.space_after = Pt(0)

    add_page_number(section.footer.paragraphs[0])

    snippets = [
        {
            "title": "代码 1：七路循迹信号采集与偏差计算功能实现",
            "file": "line_tracker.c",
            "start": 50,
            "end": 135,
            "explanation": (
                "该代码采用七通道位置权值对有效传感器进行加权平均，得到车辆相对赛道中心的横向偏差；"
                "同时通过位掩码识别左右直角特征，为后续转向控制提供方向和偏差输入。"
            ),
        },
        {
            "title": "代码 2：电机速度 PID 与循迹 PID 功能实现",
            "file": "car_control.c",
            "start": 582,
            "end": 620,
            "explanation": (
                "电机速度环根据目标编码器计数与实测计数形成闭环调速；循迹环在比例、积分和微分项基础上引入陀螺仪阻尼，"
                "用于抑制车辆偏航振荡。两级输出均进行积分与输出限幅，以降低积分饱和和执行器过驱动风险。"
            ),
        },
        {
            "title": "代码 3：比赛运行平滑加减速规划功能实现",
            "file": "competition_tasks.c",
            "start": 611,
            "end": 663,
            "explanation": (
                "该代码依据运行时间和里程计圈程进度生成速度指令：起步阶段采用 SmoothStep 曲线平滑加速，"
                "接近终点时按标定进度平滑减速至最终速度，从而兼顾运行效率、循迹稳定性和停车精度。"
            ),
        },
    ]

    for index, snippet in enumerate(snippets):
        title = doc.add_paragraph(style=code_title)
        title.add_run(snippet["title"])

        source_lines = read_range(
            snippet["file"], snippet["start"], snippet["end"]
        )
        code_lines = wrap_selected_lines(snippet["file"], source_lines)
        code_paragraph = doc.add_paragraph(style=code_style)
        code_run = code_paragraph.add_run("\n".join(code_lines))
        code_run.font.name = "Consolas"
        code_run.font.size = Pt(10.5)
        set_east_asia_font(code_run, "Consolas")

        note = doc.add_paragraph(style=explanation)
        note.add_run("说明：" + snippet["explanation"])

        if index < len(snippets) - 1:
            note.runs[-1].add_break(WD_BREAK.PAGE)

    core_properties = doc.core_properties
    core_properties.title = "设计报告核心代码摘录"
    core_properties.subject = "MSPM0G3507 七路循迹小车核心代码"
    core_properties.author = ""
    core_properties.last_modified_by = ""
    core_properties.comments = ""
    core_properties.keywords = "MSPM0G3507; 循迹小车; PID; 速度规划"

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    create_document()

